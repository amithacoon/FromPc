import docx
from babel.dates import parse_date
from django.shortcuts import render


def shuttle_text(shuttle):
    t = ''
    for i in shuttle:
        t += i.text
    return t

def docx_replace(data):
    doc = docx.Document('static/docx/lawsuit1.docx')

    for key in data:

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, data[key])

        for p in doc.paragraphs:

            begin = 0
            for end in range(len(p.runs)):

                shuttle = p.runs[begin:end+1]

                full_text = shuttle_text(shuttle)
                if key in full_text:
                    # print('Replace：', key, '->', data[key])
                    # print([i.text for i in shuttle])

                    # find the beginning
                    index = full_text.index(key)
                    # print('full_text length', len(full_text), 'index:', index)
                    while index >= len(p.runs[begin].text):
                        index -= len(p.runs[begin].text)
                        begin += 1

                    shuttle = p.runs[begin:end+1]

                    # do replace
                    # print('before replace', [i.text for i in shuttle])
                    if key in shuttle[0].text:
                        shuttle[0].text = shuttle[0].text.replace(key, data[key])
                    else:
                        replace_begin_index = shuttle_text(shuttle).index(key)
                        replace_end_index = replace_begin_index + len(key)
                        replace_end_index_in_last_run = replace_end_index - len(shuttle_text(shuttle[:-1]))
                        shuttle[0].text = shuttle[0].text[:replace_begin_index] + data[key]

                        # clear middle runs
                        for i in shuttle[1:-1]:
                            i.text = ''

                        # keep last run
                        shuttle[-1].text = shuttle[-1].text[replace_end_index_in_last_run:]

                    print('after replace', [i.text for i in shuttle])

                    # set begin to next
                    begin = end
    doc.save('static/docx/destination.docx')


def insert_infoBuilder(request):
    if request.method == 'POST':
        city = request.POST.get('city')
    full_name = request.POST.get('full_name')
    Userid = request.POST.get('Userid')
    full_address = request.POST.get('full_address')
    userphone = request.POST.get('userphone')
    companyname = request.POST.get('companyname')
    companyid = request.POST.get('companyid')
    companyfulladdress = request.POST.get('companyfulladdress')
    companycity = request.POST.get('companycity')
    companyphone = request.POST.get('companyphone')
    typeofwork = request.POST.get('typeofwork')
    netba = request.POST.get('netba')
    netbaid = request.POST.get('netbaid')
    netbaadress = request.POST.get('netbaadress')
    netbaphone = request.POST.get('netbaphone')
    num_messages = request.POST.get('num_messages')
    netbaphone = request.POST.get('netbaphone')
    hourofmessage = request.POST.get('hourofmessage_1')
    dateofmessage = request.POST.get('dateofmessage_1')
    whatdidyoudo = request.POST.get('whatdidyoudo')
    companyresponse = request.POST.get('companyresponse')
    if request.POST['unsubscribed'] == 'no':
        unsubscribe= "יודגש כי הנתבע לא סיפק אפשרות חוקית להסרת התובע מקבלת הודעות פרסומיות, ובלאו הכי לא קיבל לכך .אישור מראש מהתובע"
    else: unsubscribe=''


    print(hourofmessage)
    print(dateofmessage)

    num_messages = int(request.POST['num_messages'])
    dates_and_times = []
    date_of_message = [''] * num_messages
    time_of_message = [''] * num_messages
    for i in range(num_messages):
        date_of_message[i] = parse_date(request.POST['message_{}_date'.format(i)]).strftime("%d/%m/%Y")
        time_of_message[i] = (request.POST['message_{}_time'.format(i)])
        dates_and_times.append((date_of_message[i], time_of_message[i]))
    date_of_message = ""
    time_of_message = ""
    messages = ""
    for date, time in dates_and_times[:1:1]:
        messages += f"ביום {date}  בשעה {time}\n"
    for date, time in dates_and_times[1::1]:
        messages += f"ביום {date}  בשעה {time}\n"
    messages = messages[:-1]
    num_messages= str(num_messages)
    print(dates_and_times)  # Output:

    my_dict = {'courtcity': city,
               'full_name': full_name,
               'userid': Userid,
               'full_address': full_address,
               'userphone': userphone,
               'companyname' : companyname,
               'companyid': companyid,
               'companyfulladdress': companyfulladdress,
               'companycity': companycity,
               'companyphone': companyphone,
               'typeofwork': typeofwork,
               'netba1': netba,
               'netbaid': netbaid,
               'netbaadress': netbaadress,
               'netbaphone': netbaphone,
               'unsubscribe':unsubscribe,
               'טיפיש': num_messages,
               'ffff': num_messages,
               # 'hourofmessage': hourofmessagesmessage,
               # 'dateofmessage': messages,
               'whatdidyoudo': whatdidyoudo,
               'companyrespnse': companyresponse,
               }

    keywords = {
        'dateofmessage': messages,

    }
    docx_replace(my_dict)
    document = docx.Document('static/docx/destination.docx')

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for key, value in keywords.items():
                run.text = run.text.replace(key, value)
    document.save('static/docx/destination.docx')

    download_success = True
    filetype = 'lawsuit'
    return render(request, 'form.html', {'download_success': download_success, 'type': filetype})



