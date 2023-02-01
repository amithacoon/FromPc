from django.utils.dateparse import parse_date
from datetime import datetime
from django.http import HttpResponse

import os
from django.conf import settings
from django.http import JsonResponse
import datetime
import docx
from django.shortcuts import render, redirect
from docx.shared import Inches
from docx2pdf import convert
from djangoProject import settings

import pdfkit


def index(request):
    return render(request, 'index.html')


def letter(request):
    if request.method == 'POST':
        print(request.POST)
    return render(request, 'letter.html')

def lawsuit(request):
    return render(request, 'lawsuits.html')

def why(request):
    return render(request, 'why.html')


LIBRE_OFFICE = r"/usr/bin/soffice"

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()






current_upload = 0


def file_upload(request):
    global current_upload
    if request.method == 'POST':
        numMessages = int(request.POST.get('numMessages'))
        file = request.FILES.get('file')
        if file:  # check if file exists
            file_name = f'temp{current_upload}.png'
            file_path = os.path.join(settings.MEDIA_ROOT, file_name)
            with open(file_path, 'wb+') as destination:
                for chunk in file.chunks():
                    destination.write(chunk)
            current_upload += 1
            if current_upload == numMessages:
                current_upload = 0
                return JsonResponse({'status': 'success', 'message': 'All files uploaded successfully'})
            else:
                return JsonResponse({'status': 'success', 'message': 'File uploaded successfully'})
        else:
            return JsonResponse({'status': 'error', 'message': 'File not found'})


def letter(request):
    if request.method == 'POST':
        today = datetime.date.today()
        date_string = today.strftime("%d/%m/%Y")
        username = request.POST['full_name']
        useremail = request.POST['email']
        userphone = request.POST['phone_number']
        company_name = request.POST['company_name']
        company_address = request.POST['company_address']
        company_id = request.POST['company_id']

        if request.POST['spam_type'] == 'email':
            type_address = useremail
            received = "מייל"
            dev = "כתובת מייל"
        else:
            type_address = userphone
            received = "מסרון"
            dev = "טלפון"

        if request.POST['unsubscribed_button'] == 'false':
            could_you_unsubscribe = "אם לא די בכך, הרי שהמסרון ממילא לא עומד בדרישות החוק הצורניות בכך שאין בו אפשרות הסרה כדין"
        else:
            could_you_unsubscribe = " "

        # Get the number of message

        num_messages = int(request.POST['num_messages'])
        dates_and_times = []
        date_of_message = [''] * num_messages
        time_of_message = [''] * num_messages
        for i in range(num_messages):
            date_of_message[i] = parse_date(request.POST['message_{}_date'.format(i)]).strftime("%d/%m/%Y")
            time_of_message[i] = (request.POST['message_{}_time'.format(i)])
            dates_and_times.append((date_of_message[i], time_of_message[i]))
        date_of_message = ""
        time_of_message = ""
        messages = ""
        for date, time in dates_and_times[:1:1]:
            messages += f"ביום {date}  בשעה {time}\n"
        for date, time in dates_and_times[1::1]:
            messages += f"ביום {date}  בשעה {time}\n"

        print(dates_and_times)  # Output:

        keywords = {
            "da1te": date_string,
            "UserName": username,
            "UserEmail": useremail,
            "UserPhone": userphone,
            "dev": dev,
            "1515": type_address,
            "received": received,
            "date_of_message": messages,
            "could_you_unsubscrive": could_you_unsubscribe,
            "company_name": company_name,
            "Company_address": company_address,
            "company_id": company_id,

        }

        # Filename of the template
        template_filename = 'static/docx/template.docx'

        # Open the temp file
        document = docx.Document(template_filename)
        document.core_properties.encoding = 'UTF-8'

        # Replace keywords in the contents
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                for key, value in keywords.items():
                    run.text = run.text.replace(key, value)

        # Iterate through all the paragraphs in the docx
        hebrew_alphabet = ["א׳", "ב", "ג", "ד", "ה", "ו", "ז", "ח", "ט", "י", "כ", "ל", "מ", "נ", "ס", "ע", "פ", "צ",
                           "ק", "ר", "ש", "ת"]

        try:
            photo_path = [''] * num_messages
            # insert the photos and text for attachments
            for i in range(num_messages):
                input = ("נספח " + hebrew_alphabet[i] +
                         "\n")
                new_para = document.add_paragraph(input)
                photo_path[i] = settings.MEDIA_ROOT + ('/temp' + f'{i}' + '.png')
                new_para.add_run().add_picture(photo_path[i], width=Inches(2.5))
        except FileNotFoundError:
            return render(request, 'letter.html', {'error': 'Please attach a file before submitting the form'})


        # changes the font

        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                # Set the font to Arial
                run.font.name = "Arial"
        document.save('static/docx/modified_document.docx')
        # convert to PDF
        filetype = str(request.POST['output_file'])

        if (filetype == 'PDF'):
            # sample_doc = 'static/docx/modified_document.docx'
            # out_folder = 'static/docx'
            # convert_to_pdf(sample_doc, out_folder)
            doc = ("static/docx/modified_document.docx")
            savedoc = ("static/docx/modified_document.pdf")
            convert(doc,savedoc)

        folder = 'media'

        for filename in os.listdir(folder):
            if filename.startswith("temp") and filename.endswith(".png"):
                file_path = os.path.join(folder, filename)
                try:
                    os.unlink(file_path)
                except Exception as e:
                    print("Failed to delete file: ", e)


        download_success = True

        # document.save('static/docx/modified_document.docx')

        return render(request, 'form.html', {'download_success': download_success, 'type': filetype})

        print("Finished replacing keywords in temp file.")
    else:
        # Set the download_success flag to False
        download_success = False

        # Save the modified Word document
    return render(request, 'letter.html')


def home(request):
    return render(request, 'Home.html')


def form(request):
    # Set the download_success flag to True
    download_success = True

    # Render the form template
    return render(request, 'form.html', {'download_success': download_success})

def importError(request):

    return render(request, 'importError.html')


def download_file(request):
    if request.method == 'POST':
        if request.POST['type'] == 'PDF':
            file_path = 'static/docx/modified_document.pdf'
            filename = "Letter.pdf"
            content_type = 'application/pdf'
        if request.POST['type'] == 'lawsuit':
            file_path = 'static/docx/destination.docx'
            filename = "lawsuit.docx"
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        else:
            file_path = 'static/docx/modified_document.docx'
            filename = "Letter.docx"

            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        # Process form submission
        # Open the file in binary mode
        with open(file_path, 'rb') as f:
            # Read the file content
            file_content = f.read()

        print(file_path)
        # Create the HttpResponse object with the file content

        response = HttpResponse(file_content,
                                content_type=content_type)
        # Set the Content-Disposition header to attachment
        response['Content-Disposition'] = f'attachment; {filename}'

        # Set the download_success flag to True
        download_success = True

        # Return the HttpResponse object
        return response
    else:
        # Set the download_success flag to False
        download_success = False

    # Render the form template
    return render(request, 'form.html', {'download_success': download_success})
